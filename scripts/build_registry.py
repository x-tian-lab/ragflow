# -*- coding: utf-8 -*-
"""spec_dev 登记表构建 (D-09/D-10/D-23)。幂等:已登记且 mtime/size 未变的文件跳过。"""
import sqlite3, hashlib, os, re, sys, glob
from docx import Document

ROOT = '/sessions/peaceful-vigilant-hopper/mnt/rag sys'
JL   = '/sessions/peaceful-vigilant-hopper/mnt/metadata_output/registry.jsonl'
BATCH_SECONDS = 35

DDL = '''
CREATE TABLE IF NOT EXISTS documents (
  doc_id TEXT PRIMARY KEY,          -- sha1(relpath)
  file_path TEXT NOT NULL,          -- 相对 ROOT
  format TEXT NOT NULL,             -- docx|doc_converted|txt|xlsx
  size INTEGER, mtime REAL,
  content_hash TEXT,                -- sha256(bytes)
  doc_version TEXT,                 -- content_hash[:8]  (D-10)
  fingerprint TEXT,                 -- sha1(规范化正文)   (D-22)
  doc_title TEXT,                   -- D-23
  title_source TEXT,                -- txt_header|first_para|filename|FAIL
  effective_date TEXT,              -- D-21: 文件名后缀或txt发布日期
  source_url TEXT,
  issuing_org TEXT,
  duplicate_of TEXT,                -- D-22 填充
  version_group TEXT,               -- D-21 填充
  is_latest INTEGER,                -- D-21 填充
  status TEXT DEFAULT 'registered',
  tenant_id TEXT                    -- 预留(D-09)
);
'''

def norm_text(t):
    t = re.sub(r'\s+', '', t)
    return t

def parse_txt(path):
    raw = open(path, encoding='utf-8', errors='replace').read()
    meta = {}
    for k, field in [('标题','doc_title'),('发布日期','effective_date'),('来源','issuing_org'),('原文链接','source_url')]:
        m = re.search(rf'^{k}：(.+)$', raw[:1500], re.M)
        if m: meta[field] = m.group(1).strip()
    body = raw
    m = re.search(r'^正文：\s*$', raw, re.M)
    if m: body = raw[m.end():]
    elif 'source_url' in meta:
        m2 = re.search(re.escape(meta['source_url']), raw)
        if m2: body = raw[m2.end():]
    beff = extract_eff_date(body[:1200].splitlines())
    if beff: meta['effective_date'] = beff
    if not meta.get('doc_title'):
        for ln in body.splitlines():
            ln = ln.strip()
            if ln and len(ln) <= 60 and ln not in ('无障碍浏览',) and '下载' not in ln:
                meta['doc_title'] = ln; meta['_title_fallback'] = True
                break
    return meta, body

INVIS = re.compile(r'[\u200b\u200c\u200d\ufeff\u3000]')
def clean(s): return INVIS.sub('', s).strip()


DATE_VERB = re.compile(r'(\d{4})年(\d{1,2})月(\d{1,2})日[^\d，。]{0,40}?(第?[一二三四五六七八九十]*次?(?:修正|修订)|公布|通过|发布|施行|印发)')
def extract_eff_date(paras):
    """规则(用户裁决 2026-07-05):最晚修订/修正日期优先;无修订则用公布/通过日期"""
    revs, pubs = [], []
    for t in paras:
        for m in DATE_VERB.finditer(t):
            d = f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
            if '修' in m.group(4): revs.append(d)
            else: pubs.append(d)
    if revs: return max(revs)
    if pubs: return pubs[0]
    return None

def parse_docx(path):
    d = Document(path)
    texts = [clean(p.text) for p in d.paragraphs]
    texts = [t for t in texts if t]
    for tb in d.tables:
        for row in tb.rows:
            texts.append(' '.join(clean(c.text) for c in row.cells))
    title = texts[0] if texts else ''
    # 跨段标题合并(eval 2026-07-12 发现2增强):首段无文种后缀且不含标点 => 尝试并入下一段
    DOC_SUFFIX = ('法','决定','办法','条例','规定','通知','意见','公告','细则','规则','纲要','解释')
    if title and len(texts) > 1 and not re.search(r'[。，,.]', title) and (
            len(title) <= 10 or (len(title) <= 30 and not title.endswith(DOC_SUFFIX))):
        nxt = texts[1]
        if len(nxt) <= 40 and not re.search(r'[。]', nxt) and not nxt.startswith(('第','（','('))                 and (title + nxt).endswith(DOC_SUFFIX):
            title = title + nxt
    tsrc = 'first_para'
    if not title or len(title) > 60 or title.endswith('。'):
        tsrc = 'FAIL'
    eff = extract_eff_date(texts[:8])
    return title, tsrc, '\n'.join(texts), eff

def parse_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            parts.append('|'.join(str(v) for v in row if v is not None))
    wb.close()
    title = os.path.splitext(os.path.basename(path))[0]
    return title, 'filename', '\n'.join(parts)

def fname_date(name):
    m = re.search(r'[_（(]?(\d{8})[)）]?', name)
    if m:
        s = m.group(1)
        if s.startswith(('19','20')): return f'{s[:4]}-{s[4:6]}-{s[6:]}'
    m = re.search(r'（?(\d{4})年(\d{1,2})月', name)
    if m: return f'{m.group(1)}-{int(m.group(2)):02d}'
    return None

def collect_files():
    files = []
    for p in glob.glob(os.path.join(ROOT, '**', '*'), recursive=True):
        if not os.path.isfile(p): continue
        rel = os.path.relpath(p, ROOT)
        ext = p.lower().rsplit('.', 1)[-1] if '.' in p else ''
        if ext == 'doc':
            continue  # 原始 .doc 不入库,由 converted_docx 代表
        if ext == 'docx':
            fmt = 'doc_converted' if rel.startswith('converted_docx') else 'docx'
        elif ext == 'txt': fmt = 'txt'
        elif ext == 'xlsx': fmt = 'xlsx'
        elif ext == 'xls': continue  # 已转 xlsx
        else: continue
        files.append((rel, fmt))
    return files

def main():
    import time, json
    t0 = time.time()
    known = {}
    if os.path.exists(JL):
        for line in open(JL, encoding='utf-8'):
            r = json.loads(line)
            known[r['doc_id']] = (r['size'], r['mtime'])
    files = collect_files()
    out = open(JL, 'a', encoding='utf-8')
    done = skip = fail = 0
    for rel, fmt in files:
        if time.time() - t0 > BATCH_SECONDS: break
        p = os.path.join(ROOT, rel)
        st = os.stat(p)
        doc_id = hashlib.sha1(rel.encode()).hexdigest()[:16]
        if doc_id in known and known[doc_id] == (st.st_size, st.st_mtime):
            skip += 1; continue
        rec = {'doc_id':doc_id,'file_path':rel,'format':fmt,'size':st.st_size,'mtime':st.st_mtime,
               'status':'registered'}
        try:
            raw = open(p, 'rb').read()
            chash = hashlib.sha256(raw).hexdigest()
            meta = {'doc_title':'','title_source':'FAIL','effective_date':None,'source_url':None,'issuing_org':None}
            if fmt == 'txt':
                m, body = parse_txt(p); meta.update(m)
                if meta.pop('_title_fallback', None): meta['title_source'] = 'body_first_line'
                elif meta['doc_title']: meta['title_source'] = 'txt_header'
            elif fmt in ('docx','doc_converted'):
                t, tsrc, body, eff = parse_docx(p); meta['doc_title'], meta['title_source'] = t, tsrc
                fd = fname_date(os.path.basename(rel))
                # 规则:最晚修订>文件名>公布;extract 已保证修订优先,两者都有取较大
                meta['effective_date'] = max([d for d in (eff, fd) if d], default=None)
            else:
                t, tsrc, body = parse_xlsx(p); meta['doc_title'], meta['title_source'] = t, tsrc
            if not meta.get('effective_date'):
                meta['effective_date'] = fname_date(os.path.basename(rel))
            rec.update(meta)
            rec['content_hash'] = chash; rec['doc_version'] = chash[:8]
            rec['fingerprint'] = hashlib.sha1(norm_text(body).encode()).hexdigest()[:16] if body.strip() else None
            done += 1
        except Exception as e:
            rec['status'] = f'ERROR:{str(e)[:80]}'; fail += 1
        out.write(json.dumps(rec, ensure_ascii=False) + '\n')
        known[doc_id] = (st.st_size, st.st_mtime)
    out.close()
    print(f'本批 {done} 跳过 {skip} 失败 {fail} | 已登记 {len(known)}/{len(files)}')
    print('REMAINING' if len(known) < len(files) else 'ALL_DONE')

if __name__ == '__main__':
    main()
